import gradio as gr  
import pandas as pd  
import os  
import datetime  
from fpdf import FPDF  
import docx  
import tempfile  
import io  
import csv  
import zipfile  
  
# -----------------------------  
# Config  
# -----------------------------  
SUPPORTED_DOMAINS = ["AE", "DM", "VS", "LB", "EX", "CM"]  
EXPECTED_COLUMNS = {  
    "AE": ["STUDYID", "USUBJID", "AESEQ", "AESTDTC", "AEENDTC", "AEDECOD", "AESEV"],  
    "DM": ["STUDYID", "USUBJID", "BRTHDTC", "SEX", "RACE"],  
    "VS": ["STUDYID", "USUBJID", "VSTEST", "VSTRESN", "VSDTC"],  
}  
  
# -----------------------------  
# Cleaning & validation  
# -----------------------------  
def clean_dataframe(df: pd.DataFrame):
    df = df.copy()
    issues = {"missing": [], "invalid": [], "oor": [], "structure": [], "duplicates": []}

    for col in df.columns:
        if df[col].isna().any():
            missing_idx = df[df[col].isna()].index.tolist()
            issues["missing"].append(f"{col}: rows {missing_idx}")
            if not pd.api.types.is_numeric_dtype(df[col]):
                df[col] = df[col].fillna("UNKNOWN")

        coerced = pd.to_numeric(df[col], errors="coerce")
        if coerced.notna().any() and coerced.isna().any():
            invalid_idx = df[coerced.isna()].index.tolist()
            issues["invalid"].append(f"{col}: non-numeric values at rows {invalid_idx}")

    col_age = next((c for c in df.columns if c.upper() == "AGE"), None)
    if col_age:
        age_num = pd.to_numeric(df[col_age], errors="coerce")
        valid_age = age_num[(age_num >= 0) & (age_num <= 120)]
        avg_age = valid_age.mean() if len(valid_age) > 0 else None
        low_idx = df[age_num < 0].index.tolist()
        high_idx = df[age_num > 120].index.tolist()
        if low_idx or high_idx:
            issues["oor"].append(f"{col_age}: <0 → {low_idx} | >120 → {high_idx}")
        corrected_age = age_num.copy()
        corrected_age[(age_num < 0) | (age_num > 120)] = avg_age
        df[col_age] = corrected_age.fillna(df[col_age])

    if "USUBJID" in df.columns:
        dup_mask = df.duplicated(subset=["USUBJID"], keep=False)
        if dup_mask.any():
            dup_rows = df[dup_mask].index.tolist()
            issues["duplicates"].append(f"Duplicated USUBJID at rows {dup_rows}")

    return df, issues


def validate_sdtm_structure(df: pd.DataFrame, filename: str):  
    issues = []  
    domain = None  
    base = os.path.splitext(os.path.basename(filename))[0].upper()  
    for d in EXPECTED_COLUMNS:  
        if base == d or base.startswith(d + "_") or base.endswith("_" + d) or base.startswith(d):  
            domain = d  
            break  
    if domain and domain in EXPECTED_COLUMNS:  
        expected = set(EXPECTED_COLUMNS[domain])  
        missing = expected - set(df.columns)  
        if missing:  
            issues.append(f"Missing expected columns for {domain}: {sorted(list(missing))}")  
    return domain, issues  
  
  
# -----------------------------  
# Query generator  
# -----------------------------  
def generate_queries(df: pd.DataFrame, fname: str):  
    queries = []  
    col_age = next((c for c in df.columns if c.upper() == "AGE"), None)  
    if col_age:  
        age_num = pd.to_numeric(df[col_age], errors="coerce")
        for idx, val in age_num[(age_num < 0) | (age_num > 120)].items():
            usubjid = df.loc[idx, "USUBJID"] if "USUBJID" in df.columns else f"row {idx}"
            queries.append(f"{fname}: {usubjid} → AGE={val} (out of expected range). Please verify.")  

    aecol = next((c for c in df.columns if c.upper() == "AESTDTC"), None)  
    if aecol and "BRTHDTC" in df.columns:  
        for idx, row in df.iterrows():  
            evt = pd.to_datetime(row[aecol], errors="coerce")  
            br = pd.to_datetime(row["BRTHDTC"], errors="coerce")  
            if pd.notna(evt) and pd.notna(br) and evt < br:  
                usubjid = row.get("USUBJID", f"row {idx}")  
                queries.append(f"{fname}: {usubjid} → event {aecol} < BRTHDTC. Verify dates.")  
  
    if "USUBJID" in df.columns:  
        missing_usub = df[df["USUBJID"].isna()].index.tolist()  
        if missing_usub:  
            queries.append(f"{fname}: Missing USUBJID at rows {missing_usub}.")  
  
    return queries  
  
  
# -----------------------------  
# File helpers  
# -----------------------------  
def load_files_as_dfs(file_paths):  
    if not file_paths:  
        return [], "Please upload 1–3 CSV files."  
    dfs = []  
    errors = []  
    for p in file_paths[:3]:  
        try:  
            df = pd.read_csv(p)  
            dfs.append((os.path.basename(p), df))  
        except Exception as e:  
            errors.append(f"ERROR reading {os.path.basename(p)}: {e}")  
    return dfs, ("\n".join(errors) if errors else "")  
  
  
# -----------------------------  
# Processing logic  
# -----------------------------  
def process_files(file_paths):  
    status_lines = []  
    dfs, err = load_files_as_dfs(file_paths)  
    if err:  
        status_lines.append(f"Error loading files: {err}")  
    if not dfs:  
        return "No files uploaded.", pd.DataFrame(), "\n".join(status_lines), [], []  
  
    status_lines.append(f"Loaded {len(dfs)} file(s). Beginning validation & cleaning...")  
    reports = []  
    summary_records = []  
    cleaned_state = []  
    all_queries = []  
  
    for fname, df in dfs:  
        status_lines.append(f"Processing {fname} ({df.shape[0]} rows × {df.shape[1]} cols)...")  
        domain, structure_issues = validate_sdtm_structure(df, fname)  
        if structure_issues:  
            status_lines.append(f"Structure issues for {fname}: {structure_issues}")  
  
        cleaned_df, issues = clean_dataframe(df)  
        queries = generate_queries(df, fname)  
        all_queries.extend(queries)  
  
        rep = f"### {fname} — {df.shape[0]} rows × {df.shape[1]} cols\n\n"  
        rep += f"**Domain detected:** {domain or 'Unknown'}\n\n"  
        rep += "**Structure Issues:**\n" + ("\n".join(structure_issues) if structure_issues else "None") + "\n\n"  
        rep += "**Missing Values:**\n" + ("\n".join(issues["missing"]) if issues["missing"] else "None") + "\n\n"  
        rep += "**Invalid Values:**\n" + ("\n".join(issues["invalid"]) if issues["invalid"] else "None") + "\n\n"  
        rep += "**Out-of-Range:**\n" + ("\n".join(issues["oor"]) if issues["oor"] else "None") + "\n\n"  
        rep += "**Duplicates:**\n" + ("\n".join(issues["duplicates"]) if issues["duplicates"] else "None") + "\n"  
        rep += "---\n"  
        reports.append(rep)  
  
        summary_records.append({  
            "File": fname,  
            "Rows": df.shape[0],  
            "Columns": df.shape[1],  
            "Missing_Issues": len(issues["missing"]),  
            "Invalid_Issues": len(issues["invalid"]),  
            "OOR_Issues": len(issues["oor"]),  
            "Structure_Issues": len(structure_issues),  
            "Queries_Generated": len(queries)  
        })  
  
        cleaned_state.append((fname, cleaned_df))  
        status_lines.append(f"Finished: {fname} (cleaned). Queries generated: {len(queries)}")  
  
    summary_df = pd.DataFrame(summary_records)  
    status_lines.append("Processing complete.")  
    full_report_md = "\n\n".join(reports)  
  
    return full_report_md, summary_df, "\n".join(status_lines), cleaned_state, all_queries  
  
  
def show_corrected_tables(file_paths):  
    report, summary_df, status_log, cleaned_state, queries = process_files(file_paths)  
    out1 = out2 = out3 = None  
    for i, (fname, df) in enumerate(cleaned_state[:3]):  
        if i == 0:  
            out1 = df  
        elif i == 1:  
            out2 = df  
        elif i == 2:  
            out3 = df  
    return report, summary_df, status_log, out1, out2, out3, cleaned_state, queries  
  
  
# -----------------------------  
# Reports generation  
# -----------------------------  
def download_reports(cleaned_state):  
    if not cleaned_state:  
        return None, None, None, "No cleaned data available. Run validation first."  
  
    ts = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")  
    pdf_path = os.path.join(tempfile.gettempdir(), f"clinfix_report_{ts}.pdf")  
    docx_path = os.path.join(tempfile.gettempdir(), f"clinfix_report_{ts}.docx")  
    zip_path = os.path.join(tempfile.gettempdir(), f"clinfix_csvs_{ts}.zip")  
  
    pdf = FPDF()  
    pdf.set_auto_page_break(auto=True, margin=12)  
    pdf.add_page()  
    pdf.set_font("Arial", "B", 14)  
    pdf.cell(0, 8, "ClinFix AI - Cleaned SDTM Data Report", ln=1, align="C")  
    pdf.set_font("Arial", size=10)  
    pdf.cell(0, 6, f"Generated: {datetime.datetime.now().isoformat()}", ln=1)  
    pdf.ln(4)  
    for fname, df in cleaned_state:  
        pdf.set_font("Arial", "B", 12)  
        pdf.cell(0, 6, f"{fname} ({df.shape[0]} rows × {df.shape[1]} cols)", ln=1)  
        pdf.set_font("Arial", size=8)  
        pdf.multi_cell(0, 5, df.head(20).to_string(index=False))  
        pdf.ln(4)  
    pdf.output(pdf_path)  
  
    doc = docx.Document()  
    doc.add_heading("ClinFix AI - Cleaned Data Report", level=0)  
    doc.add_paragraph(f"Generated: {datetime.datetime.now().isoformat()}")  
    for fname, df in cleaned_state:  
        doc.add_heading(fname, level=1)  
        display_df = df.head(50).fillna("").astype(str)  
        table = doc.add_table(rows=1, cols=len(display_df.columns), style="Table Grid")  
        hdr_cells = table.rows[0].cells  
        for j, col in enumerate(display_df.columns):  
            hdr_cells[j].text = str(col)  
        for r in range(display_df.shape[0]):  
            row_cells = table.add_row().cells  
            for c in range(display_df.shape[1]):  
                row_cells[c].text = display_df.iat[r, c]  
        doc.add_page_break()  
    doc.save(docx_path)  
  
    with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:  
        for fname, df in cleaned_state:  
            csv_bytes = df.to_csv(index=False).encode("utf-8")  
            zf.writestr(fname, csv_bytes)  
  
    status = f"Reports ready: PDF ({os.path.basename(pdf_path)}), DOCX ({os.path.basename(docx_path)}), ZIP ({os.path.basename(zip_path)})"  
    return pdf_path, docx_path, zip_path, status
  

# -----------------------------
# NEW SECTION — DOWNLOAD CLEANED CSV FILES
# -----------------------------
def download_clean_csvs(cleaned_state):
    if not cleaned_state:
        return None, None, None, "No cleaned data available."

    out_paths = [None, None, None]

    for i, (fname, df) in enumerate(cleaned_state[:3]):
        path = os.path.join(tempfile.gettempdir(), f"cleaned_{fname}")
        df.to_csv(path, index=False)
        out_paths[i] = path

    return out_paths[0], out_paths[1], out_paths[2], "Cleaned CSVs ready."


# -----------------------------  
# Clear & sample  
# -----------------------------  
def clear_all():  
    return None, pd.DataFrame(), "Status: Idle", None, None, None, [], "", ""  
  
def generate_sample_files():  
    tmpdir = tempfile.mkdtemp(prefix="clinfix_samples_")  
    paths = []  
    ae_path = os.path.join(tmpdir, "AE.csv")  
    ae_rows = [  
        {"STUDYID": "S1", "USUBJID": "S1-001", "AESEQ": 1, "AESTDTC": "2020-01-10", "AEDECOD": "HEADACHE", "AESEV": "MILD", "AGE": 30},  
        {"STUDYID": "S1", "USUBJID": "S1-002", "AESEQ": 1, "AESTDTC": "2019-12-01", "AEDECOD": "", "AESEV": "MODERATE", "AGE": -5},  
    ]  
    pd.DataFrame(ae_rows).to_csv(ae_path, index=False)  
    paths.append(ae_path)  
  
    dm_path = os.path.join(tmpdir, "DM.csv")  
    dm_rows = [  
        {"STUDYID": "S1", "USUBJID": "S1-001", "BRTHDTC": "1990-05-01", "SEX": "M", "RACE": "ASIAN"},  
        {"STUDYID": "S1", "USUBJID": "S1-002", "BRTHDTC": "1985-02-02", "SEX": "", "RACE": "WHITE"},  
    ]  
    pd.DataFrame(dm_rows).to_csv(dm_path, index=False)  
    paths.append(dm_path)  
  
    zip_out = os.path.join(tmpdir, "clinfix_samples.zip")  
    with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as zf:  
        for p in paths:  
            zf.write(p, arcname=os.path.basename(p))  
    return zip_out  
  
  
# -----------------------------  
# Gradio UI  
# -----------------------------  
with gr.Blocks(title="ClinFix AI — SDTM Data Cleaning Tool") as demo:  
    gr.Markdown("# ClinFix AI — SDTM Data Cleaning Tool")  
    gr.Markdown("Upload up to 3 SDTM CSVs (AE, DM, VS, LB...). The tool runs structural checks, cleans common problems, generates site queries, and produces downloadable reports.")  
  
    with gr.Row():  
        with gr.Column(scale=2):  
            file_upload = gr.Files(label="Upload CSV Files (max 3)", file_count="multiple", type="filepath")  
            with gr.Row():  
                btn_validate = gr.Button("Validate & Clean", variant="primary")  
                btn_preview = gr.Button("View Cleaned Output", variant="secondary")  
                btn_download_reports = gr.Button("Download Reports (PDF/DOCX/ZIP)", variant="primary")  
                btn_download_csv = gr.Button("Download Cleaned CSVs", variant="primary")    # <-- NEW BUTTON  
                btn_clear = gr.Button("Clear All", variant="stop")  
  
            gr.Markdown("**Quick actions**")  
            with gr.Row():  
                btn_sample = gr.Button("Create Sample Files (Demo)")  
                sample_file = gr.File(label="Download Sample ZIP")  
  
            status_md = gr.Markdown("Status: Idle")  
            log_box = gr.Textbox(label="Processing Log", value="", lines=8, interactive=False)  
  
        with gr.Column(scale=3):  
            with gr.Tabs():  
                with gr.Tab("Processing Report"):  
                    report_md = gr.Markdown("No run yet.")  
                with gr.Tab("Summary Dashboard"):  
                    summary_table = gr.Dataframe(interactive=False)  
                with gr.Tab("Cleaned Tables"):  
                    gr.Markdown("### Cleaned datasets (up to 3)")  
                    with gr.Accordion("File 1", open=True):  
                        df1 = gr.Dataframe(interactive=False)  
                    with gr.Accordion("File 2", open=False):  
                        df2 = gr.Dataframe(interactive=False)  
                    with gr.Accordion("File 3", open=False):  
                        df3 = gr.Dataframe(interactive=False)  
                with gr.Tab("Generated Queries"):  
                    queries_out = gr.Textbox(label="Auto-generated site queries", lines=10, interactive=False)  
                with gr.Tab("AI Insights"):  
                    ai_insights = gr.Markdown("No insights yet.")  
  
    pdf_file = gr.File(label="PDF Report")  
    docx_file = gr.File(label="DOCX Report")  
    zip_file = gr.File(label="Cleaned CSVs ZIP")  

    # NEW: Individual CSVs
    csv1_file = gr.File(label="Cleaned File 1 CSV")
    csv2_file = gr.File(label="Cleaned File 2 CSV")
    csv3_file = gr.File(label="Cleaned File 3 CSV")
  
    cleaned_state = gr.State([])  
    last_queries = gr.State([])  
  
    def on_create_sample():  
        path = generate_sample_files()  
        return path  
  
    btn_sample.click(on_create_sample, inputs=[], outputs=[sample_file])  
  
    def on_validate(files):  
        report, summary_df, status_log, cleaned, queries = process_files(files)  
        status = "Validation completed."  
        return report, summary_df, status, status_log, cleaned, "\n".join(queries)  
  
    btn_validate.click(  
        on_validate,  
        inputs=[file_upload],  
        outputs=[report_md, summary_table, status_md, log_box, cleaned_state, queries_out]  
    )  
  
    def on_preview(files):  
        report, summary_df, status_log, cleaned, queries = process_files(files)  
        out1 = out2 = out3 = None  
        for i, (fname, df) in enumerate(cleaned[:3]):  
            if i == 0:  
                out1 = df  
            elif i == 1:  
                out2 = df  
            elif i == 2:  
                out3 = df  
        insights = "Simple AI Insights:\n"  
        insights += f"- Files processed: {len(cleaned)}\n"  
        total_queries = len(queries)  
        insights += f"- Total auto-generated queries: {total_queries}\n"  
        if total_queries > 0:  
            insights += "- Recommendation: Review queries and prepare site queries.\n"  
        else:  
            insights += "- No automatic queries generated.\n"  
        return report, out1, out2, out3, insights  
  
    btn_preview.click(  
        on_preview,  
        inputs=[file_upload],  
        outputs=[report_md, df1, df2, df3, ai_insights]  
    )  
  
    def on_download(cleaned):  
        pdf_path, docx_path, zip_path, status = download_reports(cleaned)  
        return pdf_path, docx_path, zip_path, status  
  
    btn_download_reports.click(  
        on_download,  
        inputs=[cleaned_state],  
        outputs=[pdf_file, docx_file, zip_file, status_md]  
    )  


    # NEW — Download Cleaned CSV Files
    def on_download_csvs(cleaned):
        return download_clean_csvs(cleaned)

    btn_download_csv.click(
        on_download_csvs,
        inputs=[cleaned_state],
        outputs=[csv1_file, csv2_file, csv3_file, status_md]
    )

  
    btn_clear.click(  
        clear_all,  
        outputs=[file_upload, summary_table, status_md, df1, df2, df3, cleaned_state, queries_out, log_box]  
    )  
  
demo.launch()
